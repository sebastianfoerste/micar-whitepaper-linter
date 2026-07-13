import hashlib
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import docx
import pytest
from lxml import etree

from micar_linter.legora_workspace import (
    build_change_set,
    build_review_sidecar,
    build_workflow_pack,
    decide_change,
    lock_cell,
    reconcile_sidecar,
    render_tracked_docx,
    update_cell,
)
from micar_linter.workspace import build_whitepaper_workspace


def _workspace() -> dict:
    return build_whitepaper_workspace([Path("examples/incomplete.json")])


def test_sidecar_uses_stable_cells_and_stale_lock_conflicts() -> None:
    sidecar = build_review_sidecar(_workspace())
    locked = lock_cell(
        sidecar, sidecar["cells"][0]["target_id"], "Reviewer", 1, datetime(2026, 7, 13, tzinfo=UTC)
    )
    with pytest.raises(ValueError, match="409 Conflict"):
        lock_cell(locked, locked["cells"][0]["target_id"], "Other", 1, datetime(2026, 7, 13, tzinfo=UTC))
    changed = _workspace()
    changed["vault"]["documents"][0]["source_sha256"] = "changed"
    assert reconcile_sidecar(sidecar, changed)["stale"] is True
    updated = update_cell(
        locked,
        locked["cells"][0]["target_id"],
        "Reviewer",
        2,
        datetime(2026, 7, 13, tzinfo=UTC),
        assignee="Reviewer",
        comment="Source checked",
        decision="approved",
    )
    assert updated["cells"][0]["comments"][0]["body"] == "Source checked"
    with pytest.raises(ValueError, match="409 Conflict"):
        update_cell(
            updated,
            updated["cells"][0]["target_id"],
            "Reviewer",
            2,
            datetime(2026, 7, 13, tzinfo=UTC),
            decision="approved",
        )


def test_docx_redline_uses_tracked_insertions_and_deletions(tmp_path) -> None:
    source = tmp_path / "source.docx"
    document = docx.Document()
    document.add_paragraph("Original clause survives")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Original table survives"
    document.save(source)
    source_digest = hashlib.sha256(source.read_bytes()).hexdigest()
    change_set = build_change_set(_workspace())
    change_set["source_digests"] = {"source": source_digest}
    change_set = decide_change(change_set, change_set["changes"][0]["id"], "accepted")
    for change in list(change_set["changes"][1:]):
        change_set = decide_change(change_set, change["id"], "rejected")
    output = render_tracked_docx(source, tmp_path / "reviewed.docx", change_set)
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_digest
    with zipfile.ZipFile(output) as package:
        xml = etree.fromstring(package.read("word/document.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    assert xml.xpath("count(.//w:ins)", namespaces=ns) > 0
    assert xml.xpath("count(.//w:del)", namespaces=ns) > 0
    text = "".join(xml.xpath(".//w:t/text()", namespaces=ns))
    assert "Original clause survives" in text
    assert "Original table survives" in text
    assert xml.xpath("count(.//w:ins)", namespaces=ns) == 1
    inserted_changes = [
        "".join(node.xpath(".//w:t/text()", namespaces=ns))
        for node in xml.xpath(".//w:ins", namespaces=ns)
    ]
    assert inserted_changes == [change_set["changes"][0]["proposed_text"]]


def test_workflow_pack_blocks_unreviewed_cells() -> None:
    workspace = _workspace()
    sidecar = build_review_sidecar(workspace)
    pack = build_workflow_pack(workspace, sidecar, build_change_set(workspace))
    assert pack["run"]["status"] == "blocked"
    assert pack["run"]["external_action_allowed"] is False
