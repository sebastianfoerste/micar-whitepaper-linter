from pathlib import Path

import docx

from micar_linter.document import load_from_docx


def test_docx_parsing_edge_cases(tmp_path: Path):
    doc_path = tmp_path / "edge_cases.docx"
    doc = docx.Document()
    
    # Text with bold runs that simulate headings
    p = doc.add_paragraph()
    run = p.add_run("Risk Factors")
    run.bold = True
    doc.add_paragraph("High risk of value loss.")
    
    # Uppercase text that simulates headings
    doc.add_paragraph("ENVIRONMENTAL IMPACT")
    doc.add_paragraph("Consensus energy consumption is low.")
    
    doc.save(doc_path)
    
    sections = load_from_docx(doc_path)
    assert "risks" in sections
    assert "environmental_impact" in sections
    
    assert "High risk of value loss" in sections["risks"]
    assert "Consensus energy consumption" in sections["environmental_impact"]
